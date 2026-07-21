"""
RAG minimal sur les documents texte non structures du corpus (PRA, rapport
d'incident, notes de cellule de crise, dossier d'architecture...). Les CSV
structures ne passent jamais par ici - ils sont traites par les tools
deterministes (graph_tools, anomaly_checks, risk_calc). Le RAG sert
uniquement a retrouver un passage source quand l'information n'est que
dans du texte libre.

Choix volontaire: TF-IDF + cosine similarity (scikit-learn) plutot que des
embeddings de reseau de neurones. Aucun modele a telecharger, marche hors
ligne, rapide, et suffisant vu la taille du corpus (quelques dizaines de
paragraphes). Le contrat reste le meme: chaque passage retourne porte sa
source et son score - jamais de texte sans provenance.

Comme pour data_loader, deux sources possibles: les .docx livres dans
data/docs/ par defaut, ou des .docx uploades par l'utilisateur (n'importe
quel nom, n'importe quel contenu) pousses via `set_active_docs`.
"""
import io
from dataclasses import dataclass

import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from core.schemas import RetrievedPassage
from pathlib import Path

DOCS_DIR = Path(__file__).resolve().parents[1] / "data" / "docs"

_active_docs: dict[str, bytes] | None = None


def set_active_docs(docs: dict[str, bytes]) -> None:
    """Bascule le RAG sur des documents uploades (nom de fichier -> octets bruts)."""
    global _active_docs
    _active_docs = docs


def clear_active_docs() -> None:
    global _active_docs
    _active_docs = None


@dataclass
class _Index:
    vectorizer: TfidfVectorizer
    matrix: object
    chunks: list[str]
    sources: list[str]


def _iter_documents():
    """Cede (nom_fichier, docx.Document) depuis la source active."""
    if _active_docs is not None:
        for name, data in _active_docs.items():
            try:
                yield name, docx.Document(io.BytesIO(data))
            except Exception:
                continue  # fichier non lisible comme .docx - ignore silencieusement
    else:
        for path in sorted(DOCS_DIR.glob("*.docx")):
            yield path.name, docx.Document(path)


def _load_chunks() -> tuple[list[str], list[str]]:
    chunks, sources = [], []
    for name, d in _iter_documents():
        for para in d.paragraphs:
            text = para.text.strip()
            if len(text) >= 15:
                chunks.append(text)
                sources.append(name)
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))
                    sources.append(name)
    return chunks, sources


def build_index() -> _Index:
    chunks, sources = _load_chunks()
    vectorizer = TfidfVectorizer(strip_accents="unicode", lowercase=True)
    matrix = vectorizer.fit_transform(chunks) if chunks else None
    return _Index(vectorizer=vectorizer, matrix=matrix, chunks=chunks, sources=sources)


def search(query: str, k: int = 3) -> list[RetrievedPassage]:
    """Renvoie les k passages les plus proches de `query`, avec leur source
    et un score de similarite (0 a 1)."""
    index = build_index()
    if not index.chunks:
        return []
    query_vec = index.vectorizer.transform([query])
    scores = cosine_similarity(query_vec, index.matrix)[0]
    top_idx = scores.argsort()[::-1][:k]
    return [
        RetrievedPassage(text=index.chunks[i], source=index.sources[i], score=float(scores[i]))
        for i in top_idx if scores[i] > 0
    ]
